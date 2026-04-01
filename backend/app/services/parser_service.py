"""Document parsing service for extracting text from various file formats."""

import logging
from pathlib import Path
from typing import Literal

from app.core.exceptions import DocumentParsingError, InvalidFileTypeError

logger = logging.getLogger(__name__)

# Supported file types
FileType = Literal["pdf", "txt", "docx", "md"]


class ParserService:
    """Service for parsing and extracting text from documents."""

    SUPPORTED_EXTENSIONS: dict[str, FileType] = {
        ".pdf": "pdf",
        ".txt": "txt",
        ".docx": "docx",
        ".md": "md",
    }

    def __init__(self) -> None:
        """Initialize the parser service."""
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()

    def _check_pdf_support(self) -> bool:
        """Check if PDF parsing is available."""
        try:
            import PyPDF2  # noqa: F401

            return True
        except ImportError:
            logger.warning("PyPDF2 not installed. PDF parsing will not be available.")
            return False

    def _check_docx_support(self) -> bool:
        """Check if DOCX parsing is available."""
        try:
            import docx  # noqa: F401

            return True
        except ImportError:
            logger.warning(
                "python-docx not installed. DOCX parsing will not be available."
            )
            return False

    def get_file_type(self, file_path: Path | str) -> FileType:
        """
        Determine the file type from the file extension.

        Args:
            file_path: Path to the file.

        Returns:
            The file type as a string.

        Raises:
            InvalidFileTypeError: If the file type is not supported.
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise InvalidFileTypeError(
                file_type=extension,
                allowed_types=list(self.SUPPORTED_EXTENSIONS.keys()),
            )

        return self.SUPPORTED_EXTENSIONS[extension]

    def parse_document(
        self, file_path: Path | str, file_type: FileType | None = None
    ) -> str:
        """
        Parse a document and extract its text content.

        Args:
            file_path: Path to the document file.
            file_type: Optional file type override. If not provided, will be
                       determined from the file extension.

        Returns:
            Extracted text content from the document.

        Raises:
            DocumentParsingError: If parsing fails.
            InvalidFileTypeError: If the file type is not supported.
        """
        path = Path(file_path)

        if not path.exists():
            raise DocumentParsingError(
                message=f"File not found: {file_path}",
                details={"file_path": str(file_path)},
            )

        # Determine file type if not provided
        if file_type is None:
            file_type = self.get_file_type(path)

        logger.info(f"Parsing document: {path.name} (type: {file_type})")

        try:
            if file_type == "pdf":
                return self._parse_pdf(path)
            elif file_type == "txt":
                return self._parse_text(path)
            elif file_type == "docx":
                return self._parse_docx(path)
            elif file_type == "md":
                return self._parse_markdown(path)
            else:
                raise InvalidFileTypeError(
                    file_type=file_type,
                    allowed_types=list(self.SUPPORTED_EXTENSIONS.keys()),
                )
        except (DocumentParsingError, InvalidFileTypeError):
            raise
        except Exception as e:
            logger.exception(f"Error parsing document: {path}")
            raise DocumentParsingError(
                message=f"Failed to parse document: {str(e)}",
                details={"file_path": str(file_path), "file_type": file_type},
            ) from e

    def _parse_pdf(self, file_path: Path) -> str:
        """
        Extract text from a PDF file using PyPDF2.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Extracted text content.

        Raises:
            DocumentParsingError: If PDF parsing is not available or fails.
        """
        if not self._pdf_available:
            raise DocumentParsingError(
                message="PDF parsing is not available. Install PyPDF2.",
                details={"file_path": str(file_path)},
            )

        import PyPDF2

        text_parts: list[str] = []

        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                logger.debug(f"PDF has {num_pages} pages")

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {e}"
                        )
                        continue

        except PyPDF2.errors.PdfReadError as e:
            raise DocumentParsingError(
                message=f"Invalid or corrupted PDF file: {str(e)}",
                details={"file_path": str(file_path)},
            ) from e

        text = "\n\n".join(text_parts)

        if not text.strip():
            logger.warning(f"No text extracted from PDF: {file_path}")

        return text

    def _parse_text(self, file_path: Path) -> str:
        """
        Read text from a plain text file with encoding fallbacks.

        Args:
            file_path: Path to the text file.

        Returns:
            File content as string.

        Raises:
            DocumentParsingError: If file cannot be read.
        """
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    content = file.read()
                    logger.debug(f"Successfully read file with encoding: {encoding}")
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise DocumentParsingError(
                    message=f"Failed to read text file: {str(e)}",
                    details={"file_path": str(file_path), "encoding": encoding},
                ) from e

        # If all encodings fail, try with error handling
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                content = file.read()
                logger.warning(
                    f"Read file with replacement characters due to encoding issues: {file_path}"
                )
                return content
        except Exception as e:
            raise DocumentParsingError(
                message=f"Failed to read text file with any encoding: {str(e)}",
                details={"file_path": str(file_path)},
            ) from e

    def _parse_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.

        Raises:
            DocumentParsingError: If DOCX parsing is not available or fails.
        """
        if not self._docx_available:
            raise DocumentParsingError(
                message="DOCX parsing is not available. Install python-docx.",
                details={"file_path": str(file_path)},
            )

        import docx

        try:
            doc = docx.Document(file_path)
            paragraphs: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)

        except Exception as e:
            raise DocumentParsingError(
                message=f"Failed to parse DOCX file: {str(e)}",
                details={"file_path": str(file_path)},
            ) from e

    def _parse_markdown(self, file_path: Path) -> str:
        """
        Read markdown file content.

        Markdown files are read as plain text since the RAG system
        benefits from preserving the markdown structure.

        Args:
            file_path: Path to the markdown file.

        Returns:
            File content as string.
        """
        return self._parse_text(file_path)


# Singleton instance for dependency injection
_parser_service: ParserService | None = None


def get_parser_service() -> ParserService:
    """Get or create the parser service instance."""
    global _parser_service
    if _parser_service is None:
        _parser_service = ParserService()
    return _parser_service
